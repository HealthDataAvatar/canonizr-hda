"""Shared fixtures for gateway integration tests."""

import hashlib
import io
import os
import secrets
import time
import uuid
from collections import namedtuple
from datetime import UTC, datetime

import pytest
import requests
from azure.data.tables import TableServiceClient
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

# Gateway table names — hardcoded here deliberately.
# If these drift from gateway/app/tables.py or portal/lib/table-names.ts,
# the e2e tests in portal/tests/integration/e2e.test.ts will catch it.
GW_SUBSCRIPTIONS = "GwSubscriptions"
GW_ENCRYPTION_KEYS = "GwEncryptionKeys"
GW_API_KEYS = "GwApiKeys"
GW_JOBS = "GwJobs"
GW_USER_JOBS = "GwUserJobs"
USER_PERMISSIONS = "UserPermissions"


def pytest_collection_modifyitems(config, items):
    """When FOCUS_TESTS=1, run only tests marked @pytest.mark.focus."""
    if os.environ.get("FOCUS_TESTS") != "1":
        return
    focus_items = [item for item in items if item.get_closest_marker("focus")]
    if focus_items:
        items[:] = focus_items


EmbeddedImage = namedtuple("EmbeddedImage", ["label", "width", "height"])

GATEWAY_URL = "http://gateway:8000"
TIMEOUT = 120
POLL_INTERVAL = 0.5

# Azurite connection strings (same well-known dev credentials)
AZURITE_TABLE_CONN = os.environ.get(
    "AZURITE_TABLE_CONN",
    "DefaultEndpointsProtocol=http;AccountName=devstoreaccount1;"
    "AccountKey=Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsuFq2UVErCz4I6tq/K1SZFPTOtr/KBHBeksoGMGw==;"
    "TableEndpoint=http://azurite:10002/devstoreaccount1",
)

# Fixed test encryption key (32 zero bytes in hex)
TEST_KEY_HEX = "0" * 64


@pytest.fixture(scope="session", autouse=True)
def seed_azurite():
    """One-time setup: ensure Azurite tables and blob container exist."""
    from azure.storage.blob import BlobServiceClient

    ts = TableServiceClient.from_connection_string(AZURITE_TABLE_CONN)
    ts.create_table_if_not_exists(GW_SUBSCRIPTIONS)
    ts.create_table_if_not_exists(GW_ENCRYPTION_KEYS)
    ts.create_table_if_not_exists(GW_API_KEYS)
    ts.create_table_if_not_exists(GW_JOBS)
    ts.create_table_if_not_exists(GW_USER_JOBS)
    ts.create_table_if_not_exists(USER_PERMISSIONS)

    # Create blob container (gateway/worker need it to exist)
    blob_conn = AZURITE_TABLE_CONN.replace("TableEndpoint", "BlobEndpoint").replace(":10002/", ":10000/")
    blob_svc = BlobServiceClient.from_connection_string(blob_conn)
    try:
        blob_svc.create_container("jobs")
    except Exception:
        pass  # already exists


TestSub = namedtuple("TestSub", ["sub_id", "api_key"])


@pytest.fixture
def test_sub():
    """An isolated test subscription (unique sub_id + user_id + API key)."""
    return _seed_test_sub()


@pytest.fixture
def second_sub():
    """A second, distinct subscription — for cross-user authorization tests."""
    return _seed_test_sub()


def _seed_test_sub() -> TestSub:
    """Seed a unique sub_id + user_id + API key into Azurite. No cross-test interference."""
    suffix = uuid.uuid4().hex[:8]
    sub_id = f"test_sub_{suffix}"
    user_id = f"test_user_{suffix}"
    api_key = f"pk_{secrets.token_hex(16)}"
    key_hash = hashlib.sha256(api_key.encode()).hexdigest()

    ts = TableServiceClient.from_connection_string(AZURITE_TABLE_CONN)

    # Seed API key hash -> subscription mapping
    ts.get_table_client(GW_API_KEYS).upsert_entity(
        {
            "PartitionKey": "key",
            "RowKey": key_hash,
            "sub_id": sub_id,
            "user_id": user_id,
        }
    )

    # Seed subscription -> user mapping
    ts.get_table_client(GW_SUBSCRIPTIONS).upsert_entity(
        {
            "PartitionKey": "subscription",
            "RowKey": sub_id,
            "user_id": user_id,
            "key_name": f"key-{suffix}",
        }
    )

    # Seed encryption key
    ts.get_table_client(GW_ENCRYPTION_KEYS).upsert_entity(
        {
            "PartitionKey": "key",
            "RowKey": user_id,
            "key_hex": TEST_KEY_HEX,
        }
    )

    # Seed permissions (not blocked)
    inverted_ts = str(9_999_999_999_999 - int(time.time() * 1000)).zfill(13)
    ts.get_table_client(USER_PERMISSIONS).upsert_entity(
        {
            "PartitionKey": user_id,
            "RowKey": f"{inverted_ts}_{suffix}",
            "timestamp": datetime.now(UTC).isoformat(),
            "isAdmin": False,
            "blocked": False,
            "stripeCustomerId": "",
            "changedBy": "system",
        }
    )

    return TestSub(sub_id=sub_id, api_key=api_key)


def submit_and_poll(
    files: dict, api_key: str, headers: dict[str, str] | None = None, timeout: float = TIMEOUT
) -> tuple[requests.Response, requests.Response]:
    """Submit a file and poll until the result is ready.

    api_key is required — use the test_sub fixture for isolation.
    Raises AssertionError if the result never arrives.
    """
    merged_headers = {"Authorization": f"Bearer {api_key}", **(headers or {})}
    submit = requests.post(f"{GATEWAY_URL}/v1/canonize", files=files, headers=merged_headers, timeout=timeout)
    if submit.status_code != 202:
        pytest.fail(f"Expected 202 from POST /v1/canonize, got {submit.status_code}: {submit.text}")

    poll_url = submit.json().get("poll_url", "")
    assert poll_url, "202 response missing poll_url"

    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        result = requests.get(f"{GATEWAY_URL}{poll_url}", headers=merged_headers, timeout=timeout)
        if result.status_code != 202:
            return submit, result
        time.sleep(POLL_INTERVAL)

    pytest.fail(f"Timed out polling {poll_url} after {timeout}s")

    return submit, result


def assert_canonize_ok(data: dict) -> list[dict]:
    """Assert the response shape of a successful canonize job and return the artefacts list."""
    assert data["status"] == "ok", f"Expected status 'ok', got {data.get('status')}"
    assert "job_id" in data
    assert "metadata" in data
    meta = data["metadata"]
    assert "detected_type" in meta
    assert "input_bytes" in meta
    assert "input_hash" in meta
    assert "artefacts" in data
    artefacts = data["artefacts"]
    assert isinstance(artefacts, list)
    for a in artefacts:
        assert "name" in a
        assert "mime_type" in a
        assert "size_bytes" in a
        assert a["size_bytes"] > 0
    return artefacts


def find_artefact(artefacts: list[dict], name: str) -> dict | None:
    """Find an artefact by name in the manifest."""
    return next((a for a in artefacts if a["name"] == name), None)


def artefact_names(artefacts: list[dict]) -> set[str]:
    """Get the set of artefact names."""
    return {a["name"] for a in artefacts}


def make_png(text: str = "Hello World", width: int = 200, height: int = 100) -> bytes:
    """Generate a PNG image with text drawn on it."""
    img = Image.new("RGB", (width, height), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), text, fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def make_pdf(text: str = "This is a test PDF document.", pages: int = 1) -> bytes:
    """Generate a simple PDF with text."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for i in range(pages):
        c.drawString(72, 700, f"{text} Page {i + 1}.")
        c.showPage()
    c.save()
    return buf.getvalue()


def make_pdf_with_image(text: str = "Document with figure below.") -> bytes:
    """Generate a PDF containing text and an embedded image."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 750, text)

    img_bytes = make_png("Chart Data")
    c.drawImage(ImageReader(io.BytesIO(img_bytes)), 72, 500, width=200, height=100)

    c.showPage()
    c.save()
    return buf.getvalue()


def make_pdf_with_images(images: list[EmbeddedImage], text: str = "Document with embedded images.") -> bytes:
    """Generate a PDF containing text and multiple embedded images of varying sizes."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 800, text)

    y_cursor = 750
    for img in images:
        img_bytes = make_png(img.label, width=img.width, height=img.height)
        c.drawImage(ImageReader(io.BytesIO(img_bytes)), 72, y_cursor - img.height, width=img.width, height=img.height)
        y_cursor -= img.height + 20
        if y_cursor < 100:
            c.showPage()
            y_cursor = 750

    c.showPage()
    c.save()
    return buf.getvalue()


def make_tiff(pages: list[str]) -> bytes:
    """Generate a multi-page TIFF with text drawn on each page."""
    frames = []
    for text in pages:
        img = Image.new("RGB", (200, 100), color="white")
        draw = ImageDraw.Draw(img)
        draw.text((10, 40), text, fill="black")
        frames.append(img)
    buf = io.BytesIO()
    frames[0].save(buf, format="TIFF", save_all=True, append_images=frames[1:])
    return buf.getvalue()


def make_docx(text: str = "This is a test Word document.") -> bytes:
    """Generate a simple DOCX."""
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph(text)
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("More content here.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    """Generate a simple XLSX with a table."""
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Test"
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 10])
    ws.append(["Beta", 20])
    ws.append(["Gamma", 30])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
